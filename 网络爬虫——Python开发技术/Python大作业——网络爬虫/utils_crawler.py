"""
通用爬虫工具函数：
- 构建请求头
- 获取文章详情并保存为 Word 文档
"""

import os
import random
import re
import time
from typing import Optional

import requests
from bs4 import BeautifulSoup
import docx
from docx.shared import Pt
from docx.oxml.ns import qn

from config import USER_AGENT_LIST, ARTICLE_DIR


def get_random_headers() -> dict:
    """随机选择 User-Agent，构造请求头。"""
    headers = {"User-Agent": random.choice(USER_AGENT_LIST)}
    return headers


def ensure_article_dir(sub_folder: Optional[str] = None) -> str:
    """
    确保文章保存目录存在，返回目录路径。
    sub_folder 不为空时，在 ARTICLE_DIR 下创建对应的子文件夹。
    """
    base_dir = ARTICLE_DIR
    if sub_folder:
        base_dir = os.path.join(ARTICLE_DIR, sanitize_filename(sub_folder))
    if not os.path.exists(base_dir):
        os.makedirs(base_dir)
    return base_dir


def sanitize_filename(name: str) -> str:
    """去除文件名中的非法字符，避免保存失败。"""
    # Windows 文件名非法字符：\/:*?"<>|
    illegal_chars = '\\/:*?"<>|'
    for ch in illegal_chars:
        name = name.replace(ch, "_")
    # 避免文件名过长
    return name[:80].strip() or "文章"


CHAPTER_TITLE_PATTERN = re.compile(
    r"(第[一二三四五六七八九十百千零〇0-9]{1,4}章|^\d{1,4}章|^0\d+章|正文[一二三四五六七八九十百千零〇0-9]{1,4})"
)


def is_novel_chapter_title(title: str) -> bool:
    """判断标题是否类似小说章节（第1章、第一章、001章、正文1、第001章等）。"""
    return bool(CHAPTER_TITLE_PATTERN.search(title.strip()))


def get_article_to_word(
    url: str,
    date: Optional[str] = "",
    only_chapter: bool = False,
    sub_folder: Optional[str] = None,
) -> bool:
    """
    根据文章链接抓取内容，并保存为 Word 文档。
    解析逻辑参考示例：优先尝试 id="art_content"，再尝试 class="MsoNormal"。
    """
    headers = get_random_headers()

    def fetch_soup(page_url: str) -> BeautifulSoup:
        resp = requests.get(page_url, headers=headers, timeout=10)
        resp.raise_for_status()
        return BeautifulSoup(resp.content, "lxml")

    # 先抓取当前页
    first_soup = fetch_soup(url)

    # -------- 标题解析：尽量适配不同网站 --------
    # 优先级：.headword -> 起点小说常见结构 h1/j_chapterName -> 普通 <h1> -> <title> -> 默认
    title_text = ""
    title_tag = first_soup.select_one(".headword")
    if not title_tag:
        # 起点/笔趣阁章节标题常见结构
        title_tag = first_soup.select_one("h1") or first_soup.select_one(".j_chapterName")
    if title_tag:
        title_text = title_tag.get_text(strip=True)
    elif first_soup.title:
        title_text = first_soup.title.get_text(strip=True)
    title = title_text or "未命名文章"

    # 若启用“只下载小说章节标题”，且标题不符合章节格式，则直接跳过
    if only_chapter and not is_novel_chapter_title(title):
        return False

    # -------- 针对笔趣阁类分页：自动合并同一章节的多页 --------
    soups = [first_soup]
    article_text = first_soup.get_text()
    total_pages = None
    m = re.search(r"第\((\d+)\/(\d+)\)页", article_text)
    if m:
        # cur_page = int(m.group(1))
        total_pages = int(m.group(2))

    if total_pages and total_pages > 1:
        # 分析 URL 规则，构造其它分页地址
        # 1) 带下划线：...1293971_1.html, 1293971_2.html, ...
        m_url = re.search(r"(.+?)_(\d+)\.html$", url)
        urls: list[str] = []
        if m_url:
            base = m_url.group(1)
            for i in range(1, total_pages + 1):
                page_url = f"{base}_{i}.html"
                urls.append(page_url)
        elif url.endswith(".html"):
            # 2) 首页无下划线：...1293971.html, 1293971_2.html, ...
            base = url[:-5]
            urls.append(url)
            for i in range(2, total_pages + 1):
                urls.append(f"{base}_{i}.html")

        if urls:
            soups = []
            visited = set()
            for page_url in urls:
                if page_url in visited:
                    continue
                visited.add(page_url)
                try:
                    soups.append(fetch_soup(page_url))
                except Exception:
                    # 某一页失败时跳过，继续其它页
                    continue

    # 创建文档
    doc = docx.Document()
    # 全局中文字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    # 标题
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(18)

    # 日期与链接
    if date:
        doc.add_paragraph(date)
    doc.add_paragraph(url)

    # -------- 正文解析：尝试多种常见布局（对所有分页依次追加） --------
    selectors = [
        "#art_content",        # 参考原示例
        ".MsoNormal",          # 参考原示例
        ".read-content",       # 起点等小说站常见
        ".read-content p",
        "#content",
        "article",
        "article p",
        ".content",
        ".content p",
    ]

    has_content = False

    for soup in soups:
        content_blocks = []
        for sel in selectors:
            blocks = soup.select(sel)
            if blocks:
                content_blocks.extend(blocks)

        # 如果以上选择器都没有匹配到，则退而求其次抓取所有 <p>
        if not content_blocks:
            content_blocks = soup.find_all("p")

        # 正文：尽量按照网页中的段落 / 换行结构来分段
        if content_blocks:
            for block in content_blocks:
                # 将 <br> 等换行标签显式替换为换行符
                for br in block.find_all("br"):
                    br.replace_with("\n")

                raw_text = block.get_text("\n")  # 用换行符作为不同子节点之间的分隔
                for line in raw_text.splitlines():
                    line = line.strip()
                    # 跳过特别短或空的行，避免太碎
                    if not line:
                        continue
                    doc.add_paragraph(line)
                    has_content = True

    if not has_content:
        # 如果未能解析到正文，则至少写入提示
        doc.add_paragraph("（未能自动解析正文，请手动查看原网页）")

    save_dir = ensure_article_dir(sub_folder=sub_folder)
    base_name = sanitize_filename(title)
    filename = base_name + ".docx"
    save_path = os.path.join(save_dir, filename)

    # 若遇到权限问题（如文件被占用/同名冲突），自动换一个新文件名重试
    attempt = 0
    while True:
        try:
            doc.save(save_path)
            break
        except PermissionError:
            attempt += 1
            ts = int(time.time())
            filename = f"{base_name}_{ts}_{attempt}.docx"
            save_path = os.path.join(save_dir, filename)

    return True


